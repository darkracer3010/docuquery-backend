from docx import Document
from typing import List
from app.documents.parsers.base import BaseParser
from app.documents.schemas import ParsedElement
import io


class DocxParser(BaseParser):
    """Parse DOCX documents using python-docx."""

    def parse(self, file_bytes: bytes, filename: str) -> List[ParsedElement]:
        elements = []
        doc = Document(io.BytesIO(file_bytes))
        current_heading = None

        for paragraph in doc.paragraphs:
            text = self._clean_text(paragraph.text)
            if not text or len(text) < 3:
                continue

            style_name = paragraph.style.name.lower() if paragraph.style else ""

            if "heading" in style_name:
                current_heading = text
                elements.append(
                    ParsedElement(
                        text=text,
                        element_type="heading",
                        metadata={
                            "source": filename,
                            "heading_level": style_name,
                        },
                    )
                )
            else:
                elements.append(
                    ParsedElement(
                        text=text,
                        element_type="paragraph",
                        metadata={
                            "source": filename,
                            "section": current_heading or "Document",
                        },
                    )
                )

        # Also parse tables
        for table_idx, table in enumerate(doc.tables):
            headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []

            for row_idx, row in enumerate(table.rows[1:], start=1):
                row_data = {
                    headers[i] if i < len(headers) else f"col_{i}": cell.text.strip()
                    for i, cell in enumerate(row.cells)
                }
                row_text = " | ".join(f"{k}: {v}" for k, v in row_data.items() if v)
                if row_text:
                    elements.append(
                        ParsedElement(
                            text=row_text,
                            element_type="table_row",
                            metadata={
                                "source": filename,
                                "table_index": table_idx,
                                "row_index": row_idx,
                                "headers": headers,
                            },
                        )
                    )

        return elements
